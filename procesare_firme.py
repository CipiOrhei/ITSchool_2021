"""
Modul care gestioneaza firmele
"""
import glob
from docx import Document
import re

LOCATIE_FOLDER_FIRME = r"./Documente/firme_licitatie/"
LISTA_FIRME_DISPONIBILE = list()

class Firme:
    def __init__(self, nume):
        """
        Constructor pentru clasa firma
        :param nume: numele firmei
        """
        self.nume = nume
        # cream o lista cu toate locomotivele posibile
        self.lista_locomotive = list()
        self.lista_locomotive.append({'model': 'A', 'cantitate': 30, 'viteza': 100, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0})
        self.lista_locomotive.append({'model': 'B', 'cantitate': 60, 'viteza': 80, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0})
        self.lista_locomotive.append({'model': 'C', 'cantitate': 30, 'viteza': 150, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0})
        self.lista_locomotive.append({'model': 'D', 'cantitate': 50, 'viteza': 70, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0})
        self.lista_locomotive.append({'model': 'E', 'cantitate': 45, 'viteza': 75, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0})
        # cream o lista cu toate vagoanele posibile
        self.lista_vagoane = list()
        self.lista_vagoane.append({'model': 'X12', 'cantitate': 2, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Lemn', 'Busteni', 'Fier']})
        self.lista_vagoane.append({'model': 'X14', 'cantitate': 4, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Lemn', 'Busteni', 'Fier']})
        self.lista_vagoane.append({'model': 'X13', 'cantitate': 5, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Nisip', 'Cherestea', 'Piatra']})
        self.lista_vagoane.append({'model': 'X15', 'cantitate': 4, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Nisip', 'Cherestea', 'Piatra']})
        self.lista_vagoane.append({'model': 'X24', 'cantitate': 1, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Apa', 'Benzina']})
        self.lista_vagoane.append({'model': 'X23', 'cantitate': 5, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Apa', 'Benzina']})
        self.lista_vagoane.append({'model': 'X25', 'cantitate': 10, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Apa', 'Benzina']})
        self.lista_vagoane.append({'model': 'X5', 'cantitate': 1, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Pasageri']})
        self.lista_vagoane.append({'model': 'X10', 'cantitate': 5, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Pasageri']})
        self.lista_vagoane.append({'model': 'X30', 'cantitate': 10, 'unitati_disponibile': 0, 'unitati_folosite': 0, 'cost': 0, 'material': ['Pasageri']})

    def adauga_locomotiva(self, model, unitati, cost):
        for locomotiva in self.lista_locomotive:
            if model == locomotiva['model']:
                locomotiva['unitati_disponibile'] = unitati
                locomotiva['cost'] = cost
                break

    def adauga_vagon(self, model, unitati, cost):
        for vagoane in self.lista_vagoane:
            if model == vagoane['model']:
                vagoane['unitati_disponibile'] = unitati
                vagoane['cost'] = cost
                break


def citire_document_word(cale: str) -> dict:
    """
    Citeste dintr-un word format datele despre o firma
    :param cale:
    :return: Un dictionar cu datele necesare
    """
    tmp_dict = dict()
    document = Document(cale)
    print(cale)
    for para in document.paragraphs:
        # print(para.text)
        detalii_firme = para.text

        # am scos numele firmei
        nume_regex = r"(?<=Nume: ).*$"
        nume_firma = re.findall(nume_regex,detalii_firme)
        if (nume_firma):
            nume_firma = modificare_stringuri(nume_firma)
            print(f"Nume: {nume_firma}")
            tmp_dict['nume'] = nume_firma

        # am scos locomotivele
        locomotive_regex = r"(?<=Locomotive: ).*$"
        detalii_locomotive = re.findall(locomotive_regex,detalii_firme)
        if (detalii_locomotive):
            # print(detalii_locomotive)
            detalii_locomotive = modificare_stringuri(detalii_locomotive)
            # print(detalii_locomotive)
            locomotives = str(detalii_locomotive).split(",")

            locomotives_lista = list()
            for locomotive in locomotives:
                locomotive_lista = list()
                # print((locomotive).lstrip(' '))
                locomotive = ((locomotive).lstrip(' '))
                locomotive_model_regex = r"(?<=Model ).*$"
                locomotive_model = re.findall(locomotive_model_regex, locomotive)
                # print(locomotive_model)

                locomotive_cantitate_regex = r"^.*(?= x)"
                locomotive_cantitate = re.findall(locomotive_cantitate_regex, locomotive)
                # print(locomotive_cantitate)
                locomotive_cantitate = modificare_stringuri(locomotive_cantitate)
                locomotive_model = modificare_stringuri(locomotive_model)
                print(f"Locomotive cantitate: {locomotive_cantitate}")
                print(f"Locomotive model: {locomotive_model}")
                locomotive_lista.append(locomotive_cantitate)
                locomotive_lista.append(locomotive_model)

                # cost
                # print(f"Cost: {cale}")
                for para in document.paragraphs:
                    # print(para.text)
                    detalii_cost = para.text

                    cost_regex_locomotive = fr"(?<=Model {locomotive_model} ).*(?= RON/h)"
                    cost_locomotive = re.findall(cost_regex_locomotive, detalii_cost)
                    if (cost_locomotive):
                        cost_locomotive = modificare_stringuri(cost_locomotive)
                        print(f"Cost locomotive: {cost_locomotive}")
                        locomotive_lista.append(cost_locomotive)
                locomotives_lista.append(locomotive_lista)

            tmp_dict['Locomotive'] = locomotives_lista
            print(locomotives_lista)

        # am scos vagoanele
        vagoane_regex = r"(?<=Vagoane: ).*$"
        detalii_vagoane = re.findall(vagoane_regex,detalii_firme)
        if (detalii_vagoane):
            # print(detalii_vagoane)
            detalii_vagoane = modificare_stringuri(detalii_vagoane)
            # print(detalii_vagoane)
            vagons = str(detalii_vagoane).split(",")

            vagones_lista = list()
            for vagoane in vagons:
                vagoane_lista = list()
                # print((vagoane).lstrip(' '))
                vagoane = ((vagoane).lstrip(' '))
                vagoane_model_regex = r"(?<=Model ).*$"
                vagoane_model = re.findall(vagoane_model_regex,vagoane)
                # print(vagoane_model)

                vagoane_cantitate_regex = r"^.*(?= x)"
                vagoane_cantitate = re.findall(vagoane_cantitate_regex,vagoane)
                # print(vagoane_cantitate)

                vagoane_cantitate = modificare_stringuri(vagoane_cantitate)
                vagoane_model = modificare_stringuri(vagoane_model)
                print(f"Vagoane cantitate: {vagoane_cantitate}")
                print(f"Vagoane model: {vagoane_model}")
                vagoane_lista.append(vagoane_cantitate)
                vagoane_lista.append(vagoane_model)

                # cost
                # print(f"Cost: {cale}")
                for para in document.paragraphs:
                    # print(para.text)
                    detalii_cost = para.text

                    cost_regex_vagoane = fr"(?<=Model {vagoane_model} ).*(?= RON/h)"
                    cost_vagoane = re.findall(cost_regex_vagoane, detalii_cost)
                    if (cost_vagoane):
                        cost_vagoane = modificare_stringuri(cost_vagoane)
                        print(f"Cost vagoane: {cost_vagoane}")
                        vagoane_lista.append(cost_vagoane)

                vagones_lista.append(vagoane_lista)
            print(vagones_lista)
            tmp_dict['Vagoane'] = vagones_lista



        # print('TODO: cu ajutorul regex gasiti informatiile din fiecare paragraf')
        # Populam tmp_dict
        #TODO
    # exemplu:
    # tmp_dict['nume'] = 'FerovTrans SRL'
    # tupla per locomotiva: (cantitate, model, cost)
    # tmp_dict['Locomotive'] = [(5, 'C', 20), (2, 'B', 30)]
    # # tupla per vagon: (cantitate, model, cost)
    # tmp_dict['Vagoane'] = [(6, 'X30', 35), (5, 'X14', 45), (20, 'X25', 60)]
    print(tmp_dict)
    return tmp_dict

def modificare_stringuri(stringuri_de_modificare):
    char_de_schimbat = "[]'"
    for ch in char_de_schimbat:
        stringuri_de_modificare = str(stringuri_de_modificare).replace(ch, "")

    return stringuri_de_modificare



def salvam_lista_firme_pe_hard():
    #TODO salvam lista cumva :)
    #poate txt, poate numpy
    pass


def procesare_firme_inscrise():
    # aflam toate documentele depuse de firme
    lista_fisier = glob.glob(LOCATIE_FOLDER_FIRME + '/*')
    # pentru fiecare fisier gasit procesam datele
    for file in lista_fisier:
        tmp_dict = citire_document_word(file)

        LISTA_FIRME_DISPONIBILE.append(Firme(tmp_dict['nume']))
        for element in tmp_dict['Locomotive']:
            LISTA_FIRME_DISPONIBILE[-1].adauga_locomotiva(model=element[1], unitati=element[0], cost=element[2])
        for element in tmp_dict['Vagoane']:
            LISTA_FIRME_DISPONIBILE[-1].adauga_vagon(model=element[1], unitati=element[0], cost=element[2])

    salvam_lista_firme_pe_hard()


if __name__ == '__main__':
    # firma_1 = Firme(nume='Ion')
    # print(firma_1)
    procesare_firme_inscrise()
